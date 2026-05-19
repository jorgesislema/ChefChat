"""Genera documentacion profesional en Word y HTML a partir de Markdown."""

import os
import re
try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml.parser import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def markdown_to_paragraphs(texto_md: str) -> list[tuple]:
    """Parsea Markdown a una lista de tuplas (tipo, texto, nivel, subitems opcional)."""
    lineas = texto_md.split('\n')
    bloques: list[tuple] = []
    i = 0
    while i < len(lineas):
        linea = lineas[i]
        
        # Encabezados
        match_h = re.match(r'^(#{1,6})\s+(.+)', linea)
        if match_h:
            bloques.append(('h', match_h.group(2), len(match_h.group(1))))
            i += 1
            continue
        
        # Tablas (detectamos filas con |)
        if '|' in linea and '---' not in linea:
            filas_tabla = []
            es_tabla = False
            while i < len(lineas) and '|' in lineas[i]:
                fila = lineas[i]
                if '---' not in fila.replace('|', '').replace('-', '').replace(' ', ''):
                    celdas = [c.strip() for c in fila.split('|') if c.strip()]
                    if celdas:
                        filas_tabla.append(celdas)
                        es_tabla = True
                i += 1
            if es_tabla and len(filas_tabla) >= 2:
                bloques.append(('table', filas_tabla, 0))
            continue
        
        # Codigo (```)
        if linea.strip().startswith('```'):
            codigo = []
            i += 1
            while i < len(lineas) and not lineas[i].strip().startswith('```'):
                codigo.append(lineas[i])
                i += 1
            i += 1
            if codigo:
                bloques.append(('code', '\n'.join(codigo), 0))
            continue
        
        # Listas
        match_li = re.match(r'^(\s*)[-*+]\s+(.+)', linea)
        if match_li:
            texto_li = match_li.group(2)
            subitems = []
            # Buscar subitems con doble espacio
            j = i + 1
            while j < len(lineas) and re.match(r'^\s{4,}[-*+]\s+', lineas[j]):
                subitems.append(re.sub(r'^\s{4,}[-*+]\s+', '', lineas[j]))
                j += 1
            bloques.append(('li', texto_li, 0, subitems))
            i = j
            continue
        
        # Parrafos normales
        if linea.strip():
            parrafo = linea.strip()
            # Juntar lineas consecutivas del mismo parrafo
            j = i + 1
            while j < len(lineas) and lineas[j].strip() and not lineas[j].strip().startswith('#') and '|' not in lineas[j] and not lineas[j].strip().startswith('```') and not re.match(r'^[-*+]\s+', lineas[j]):
                parrafo += ' ' + lineas[j].strip()
                j += 1
            bloques.append(('p', parrafo, 0))
            i = j
            continue
        
        i += 1
    
    return bloques


def _add_paragraph_bottom_border(paragraph, val='single', sz='12', color='2563EB'):
    """Agrega un borde inferior a un párrafo."""
    # pylint: disable=protected-access
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), val)
    bottom.set(qn('w:sz'), sz)
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _render_heading(doc, texto, nivel):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
    if nivel == 1:
        run.font.size = Pt(20)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        _add_paragraph_bottom_border(p)
    elif nivel == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)


def _apply_code_cell_shading(cell, fill_color='0F172A'):
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}" w:val="clear"/>')
    # pylint: disable=protected-access
    cell._element.get_or_add_tcPr().append(shading)
    # pylint: enable=protected-access


def _render_code_block(doc, codigo):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    _apply_code_cell_shading(cell)
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(codigo)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF0)
    doc.add_paragraph()


def _render_table(doc, filas):
    if len(filas) < 2:
        return
    table = doc.add_table(rows=len(filas), cols=len(filas[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Light Grid Accent 1'
    for r_idx, fila in enumerate(filas):
        for c_idx, celda in enumerate(fila):
            cell = table.cell(r_idx, c_idx)
            cell.text = celda
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if r_idx == 0:
                        run.bold = True
    doc.add_paragraph()


def _render_list_item(doc, texto, subitems):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texto)
    run.font.size = Pt(10)
    for sub in subitems:
        p2 = doc.add_paragraph(style='List Bullet 2')
        run2 = p2.add_run(sub)
        run2.font.size = Pt(9)


def _render_paragraph(doc, texto):
    p = doc.add_paragraph()
    partes = re.split(r'(\*\*.+?\*\*|\*.+?\*|`.+?`)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = p.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith('*') and parte.endswith('*'):
            run = p.add_run(parte[1:-1])
            run.italic = True
        elif parte.startswith('`') and parte.endswith('`'):
            run = p.add_run(parte[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
        else:
            p.add_run(parte)
    p.paragraph_format.space_after = Pt(6)


def generar_docx(ruta_md: str, ruta_docx: str):
    """Genera un documento Word profesional desde Markdown."""
    with open(ruta_md, "r", encoding="utf-8") as f:
        texto_md = f.read()
    
    bloques = markdown_to_paragraphs(texto_md)
    doc = Document()
    
    # --- Estilos ---
    # pylint: disable=no-member
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
    style.paragraph_format.space_after = Pt(6)
    # pylint: enable=no-member
    
    # --- Margenes ---
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
    
    # --- Portada ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\n\n\n\n')
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ChefChat Pro v1.0')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x25, 0x63, 0xEB)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Manual Operativo y Arquitectura de Sistema')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x47, 0x53, 0x69)
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('\nDocumentacion Confidencial - Solo para uso interno')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True
    
    doc.add_page_break()
    
    # --- Procesar bloques ---
    for bloque in bloques:
        tipo = bloque[0]
        
        if tipo == 'h':
            _render_heading(doc, bloque[1], bloque[2])
        elif tipo == 'table':
            _render_table(doc, bloque[1])
        elif tipo == 'code':
            _render_code_block(doc, bloque[1])
        elif tipo == 'li':
            texto = bloque[1]
            subitems = bloque[3] if len(bloque) > 3 else []
            _render_list_item(doc, texto, subitems)
        elif tipo == 'p':
            _render_paragraph(doc, bloque[1])
    
    # --- Footer ---
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('--- Fin del Documento ---')
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    run.italic = True
    
    doc.save(ruta_docx)
    return True


def generar_html(ruta_md: str, ruta_html: str):
    """Genera HTML profesional desde Markdown (sin dependencias externas)."""
    import markdown
    with open(ruta_md, "r", encoding="utf-8") as f:
        texto_md = f.read()
    
    html_body = markdown.markdown(texto_md, extensions=['tables', 'fenced_code', 'codehilite'])
    
    html = f"""<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <title>ChefChat Pro v1.0 - Documentacion</title>
    <style>
        @page {{ size: A4; margin: 20mm 15mm; }}
        body {{ font-family: 'Segoe UI', Arial, sans-serif; color: #1E293B; line-height: 1.7; font-size: 10pt; max-width: 800px; margin: 0 auto; padding: 20px; }}
        h1 {{ color: #2563EB; border-bottom: 3px solid #2563EB; padding-bottom: 10px; }}
        h2 {{ color: #1E3A8A; border-left: 4px solid #2563EB; padding-left: 10px; margin-top: 25px; }}
        h3 {{ color: #334155; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th {{ background: #1E293B; color: #FFF; padding: 8px; text-align: left; }}
        td {{ border: 1px solid #E2E8F0; padding: 8px; }}
        tr:nth-child(even) {{ background: #F8FAFC; }}
        code {{ background: #E2E8F0; padding: 2px 5px; border-radius: 3px; font-family: 'Consolas', monospace; font-size: 9pt; }}
        pre {{ background: #1E293B; color: #F8FAFC; padding: 15px; border-radius: 5px; overflow-x: auto; font-size: 9pt; }}
        @media print {{ body {{ max-width: none; }} }}
    </style>
</head>
<body>
    <h1>ChefChat Pro v1.0 - Manual Operativo</h1>
    {html_body}
</body>
</html>"""
    
    with open(ruta_html, "w", encoding="utf-8") as f:
        f.write(html)
    return True


def generar_documentacion():
    ruta_md = "docs_internos/resumen.md"
    
    if not os.path.exists(ruta_md):
        print(f"ERROR: No se encuentra el archivo {ruta_md}")
        return
    
    if HAS_DOCX:
        ruta_docx = os.path.abspath("docs_internos/ChefChat_SOP_Documentacion.docx")
        print("Generando documento Word profesional...")
        if generar_docx(ruta_md, ruta_docx):
            print(f"Documento Word generado: {ruta_docx}")
    else:
        print("python-docx no instalado. Instala con: pip install python-docx")
    
    ruta_html = "docs_internos/ChefChat_SOP_Documentacion.html"
    print("Generando version HTML...")
    if generar_html(ruta_md, ruta_html):
        print(f"HTML generado: {ruta_html}")
        abs_html = os.path.abspath(ruta_html)
        if os.path.exists(abs_html):
            os.startfile(abs_html)
            print(f"Abriendo en navegador...")


if __name__ == "__main__":
    generar_documentacion()
